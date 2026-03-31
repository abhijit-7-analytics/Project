"""
SalesDB — Updated 75-85 Page Report Generator
Adds: SDLC Chapter, Source Code Chapter, Expanded Content
Run:   python report_final.py
Output: Retail_Sales_Data_Analysis_Final.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

FONT = 'Arial'
H1 = 14   # Chapter heading
H2 = 12   # Section heading
H3 = 12   # Subsection heading
BODY = 11  # Body text
SMALL = 10
CODE_FONT = 'Consolas'
CODE_SZ = 8

# ═══ STUDENT INFO ═══
PROJECT = "Retail Sales Data Analysis"
STUDENT = "Barsha Priyadarsini Parida"
ROLL = "24MC034"
GUIDE = "Prof. Prangya Paramita Mohapatra"
GUIDE_TITLE = "Associate Professor, Dept. of MCA"
COLLEGE = "GITA AUTONOMOUS COLLEGE"
ADDR = "BHUBANESWAR — 752054"
SESSION = "2024-26"
DEGREE = "Master in Computer Application"


# ═══ FORMATTING HELPERS ═══

def sf(run, name=FONT, size=BODY, bold=False, italic=False,
       color=None, underline=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)


def para(doc, text, size=BODY, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, sa=6, sb=0,
         color=None, indent=None, line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(sa)
    pf.space_before = Pt(sb)
    pf.line_spacing = line_spacing
    if indent:
        pf.left_indent = Cm(indent)
    run = p.add_run(text)
    sf(run, FONT, size, bold, italic, color)
    return p


def heading(doc, text, level=1, number=""):
    sizes = {1: H1, 2: H2, 3: H3, 4: BODY}
    sz = sizes.get(level, BODY)
    full = f"{number}  {text}" if number else text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 12)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.5
    run = p.add_run(full.upper() if level == 1 else full)
    sf(run, FONT, sz, bold=True)
    return p


def chapter_title(doc, chapter_num, title):
    para(doc, "", sa=30)
    para(doc, f"CHAPTER {chapter_num}", H1 + 2, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    para(doc, title.upper(), H1 + 2, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)


def bullet(doc, text, indent=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"•  {text}")
    sf(run, FONT, BODY)
    return p


def numbered(doc, num, text, indent=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"{num}. ")
    sf(run, FONT, BODY, bold=True)
    run = p.add_run(text)
    sf(run, FONT, BODY)
    return p


def code_block(doc, code, caption=""):
    if caption:
        para(doc, caption, SMALL, italic=True,
             align=WD_ALIGN_PARAGRAPH.LEFT, sa=2)
    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.0
        pf.left_indent = Cm(0.5)
        run = p.add_run(line)
        sf(run, CODE_FONT, CODE_SZ)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
        p._element.get_or_add_pPr().append(shading)
    para(doc, "", sa=4)


def add_table(doc, headers, rows, caption=""):
    if caption:
        para(doc, caption, SMALL, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        sf(run, FONT, SMALL, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2563EB"/>')
        cell._element.get_or_add_tcPr().append(shading)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            sf(run, FONT, SMALL)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r % 2 == 0:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="F0F4FF"/>')
                cell._element.get_or_add_tcPr().append(shading)
    para(doc, "", sa=6)


def page_break(doc):
    doc.add_page_break()


def blank(doc, n=1):
    for _ in range(n):
        para(doc, "", sa=0)


def header_footer(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)
        h = section.header
        hp = h.paragraphs[0] if h.paragraphs else h.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = hp.add_run(f"{PROJECT} | Project Report")
        sf(run, FONT, 8, italic=True, color=(150, 150, 150))
        f = section.footer
        fp = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = fp.add_run()
        fld = (f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "'
               f'><w:r><w:t>1</w:t></w:r></w:fldSimple>')
        run._element.append(parse_xml(fld))
        sf(run, FONT, 9)


# ═══════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════

def cover_page(doc):
    blank(doc, 2)
    para(doc, DEGREE.upper(), H1, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
    para(doc, "Project Report on", BODY, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(PROJECT)
    sf(run, FONT, H1 + 2, bold=True)
    run.font.underline = True
    run.font.color.rgb = RGBColor(26, 26, 46)
    blank(doc, 2)
    para(doc, "Under the Able Guidance of", H2, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    para(doc, GUIDE, H2, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para(doc, GUIDE_TITLE, BODY, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
    blank(doc, 1)
    para(doc, "Submitted By", H2, bold=True, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=4,
         color=(180, 0, 0))
    para(doc, STUDENT, H2, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para(doc, f"Roll No.: {ROLL}", BODY, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
    blank(doc, 2)
    para(doc, "[COLLEGE LOGO]", BODY, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER,
         color=(180, 180, 180), sa=8)
    para(doc, f"DEPARTMENT OF {DEGREE.upper()}", H2, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para(doc, COLLEGE, H2, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para(doc, ADDR, BODY,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para(doc, SESSION, H2, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)


# ═══════════════════════════════════════════
#  CERTIFICATE, DECLARATION, ACKNOWLEDGEMENT
# ═══════════════════════════════════════════

def certificate(doc):
    para(doc, COLLEGE, H1, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2, sb=20)
    para(doc, ADDR, BODY,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
    para(doc, "[COLLEGE LOGO]", BODY, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER,
         color=(180, 180, 180), sa=16)
    heading(doc, "Certificate", 1)
    blank(doc, 1)
    para(doc, (
        f'This is to certify that the project entitled "{PROJECT}" '
        f'being submitted by Ms. {STUDENT} bearing Registration No. '
        f'{ROLL} in partial fulfilment of the requirement for the award '
        f'of the degree of {DEGREE} is a bonafide work carried out at '
        f'{COLLEGE} under my Supervision.'
    ))
    blank(doc, 5)
    t = doc.add_table(rows=1, cols=3)
    for i, label in enumerate(["Signature of Student",
                                "Project Guide",
                                "Head of Department"]):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(label)
        sf(run, FONT, BODY, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_break(doc)


def declaration(doc):
    para(doc, COLLEGE, H1, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2, sb=20)
    para(doc, ADDR, BODY,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
    heading(doc, "Declaration", 1)
    blank(doc, 1)
    para(doc, (
        f"I hereby declare that the matter embodied in this project "
        f"report is original and has not been submitted for the award "
        f"of any other degree."
    ))
    blank(doc, 5)
    para(doc, "Signature", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"Name: {STUDENT}",
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"Regd. No.: {ROLL}",
         align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)


def acknowledgement(doc):
    para(doc, COLLEGE, H1, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2, sb=20)
    para(doc, ADDR, BODY,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
    heading(doc, "Acknowledgement", 1)
    blank(doc, 1)
    para(doc, (
        f"With immense pleasure I, Ms. {STUDENT}, presenting "
        f'"{PROJECT}" project report as part of the curriculum of MCA. '
        f"I would like to express my special thanks of gratitude to "
        f"{GUIDE} for able guidance and support in completing my project."
    ))
    para(doc, (
        f"I express my profound thanks to Head of The Department "
        f"for moral support and guidance. And all those who have "
        f"indirectly guided and helped me in preparation of this project."
    ))
    blank(doc, 3)
    para(doc, "Signature of Student",
         align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)


# ═══════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════

def abstract(doc):
    heading(doc, "ABSTRACT", 1)
    blank(doc, 1)
    para(doc, (
        f'The present project, "{PROJECT}", is a comprehensive '
        f'full-stack web application developed as the major project for '
        f'the degree of {DEGREE} at {COLLEGE}, Bhubaneswar. The project '
        f'addresses the real-world problem of fragmented and manual sales '
        f'record management in small and medium-sized retail enterprises '
        f'by providing a unified, automated, and analytically rich '
        f'platform that covers the complete retail sales lifecycle [1].'
    ))
    para(doc, (
        "The system is implemented using Flask (Python) as the backend "
        "web framework, PostgreSQL as the relational database management "
        "system, and plain HTML5, CSS3, and JavaScript as the frontend "
        "technologies. The backend exposes a RESTful API consisting of "
        "more than twenty endpoints covering authentication, customer "
        "management, product management, sales operations, stock "
        "tracking, invoice generation, and report analytics. Chart.js is "
        "used for interactive data visualisation, and ReportLab is used "
        "for server-side PDF invoice generation [2]."
    ))
    para(doc, (
        "The application provides a dark-themed, responsive dashboard "
        "displaying eight live KPI metrics including total revenue "
        "($2,120.00), total sales count (6), average order value "
        "($353.33), customer count (6), product count (8), low stock "
        "items (3), out-of-stock items (2), and pending invoice count "
        "(3). Revenue distribution across product categories "
        "(Electronics, Home Goods, Apparel, Other) is visualised through "
        "bar charts and doughnut charts. Customer segmentation between "
        "Gold and Regular membership types is presented via pie charts [3]."
    ))
    para(doc, (
        "The invoice module features both automatic invoice creation "
        "upon sale completion and a standalone Invoice Builder that "
        "auto-populates customer sales data on customer ID selection, "
        "supports manual line item addition, computes GST at 18%, "
        "generates live preview modals, and downloads professional PDF "
        "invoices with company branding. All write operations are "
        "protected by session-based admin authentication with a "
        "30-minute automatic expiry."
    ))
    para(doc, (
        "The stock management module tracks product-wise stock levels "
        "with colour-coded status indicators (In Stock, Low Stock, Out "
        "of Stock), provides a restocking interface, and maintains a "
        "complete stock change history log. The reports module generates "
        "filterable sales analytics by date range, product category, and "
        "region, with CSV export capability."
    ))
    para(doc, (
        "The project follows the Waterfall Model of the Software "
        "Development Life Cycle (SDLC), progressing through five "
        "sequential phases: Requirement Analysis, System Design, "
        "Implementation (Coding), Testing, and Deployment & Maintenance. "
        "This report documents the complete SDLC process across ten "
        "chapters, from requirements analysis through design, "
        "implementation, testing, and evaluation."
    ))
    para(doc, (
        "Keywords: Retail Sales Analytics, Flask, PostgreSQL, RESTful "
        "API, Invoice Generation, Stock Management, Chart.js, ReportLab, "
        "GST, Dashboard KPI, Waterfall Model, SDLC."
    ), italic=True, sb=8)
    page_break(doc)


# ═══════════════════════════════════════════
#  TABLE OF CONTENTS (Updated with SDLC + Code)
# ═══════════════════════════════════════════

def table_of_contents(doc):
    heading(doc, "TABLE OF CONTENTS", 1)
    blank(doc, 1)
    toc = [
        ("", "Declaration", "i"),
        ("", "Certificate", "ii"),
        ("", "Acknowledgement", "iii"),
        ("", "Abstract", "iv"),
        ("", "Table of Contents", "v"),
        ("", "List of Figures", "vii"),
        ("", "List of Tables", "viii"),
        ("Chapter I", "INTRODUCTION", "1"),
        ("1.1", "Background and Motivation", "1"),
        ("1.2", "Problem Statement", "3"),
        ("1.3", "Objectives of the Project", "4"),
        ("1.4", "Scope of the Project", "5"),
        ("1.5", "Organisation of the Report", "6"),
        ("Chapter II", "LITERATURE REVIEW", "8"),
        ("2.1", "Overview of Web-Based Sales Systems", "8"),
        ("2.2", "Related Work and Existing Systems", "10"),
        ("2.3", "Technologies Reviewed", "12"),
        ("2.4", "Summary and Research Gap", "14"),
        ("Chapter III", "SYSTEM ANALYSIS AND REQUIREMENTS", "16"),
        ("3.1", "Feasibility Study", "16"),
        ("3.2", "Functional Requirements", "18"),
        ("3.3", "Non-Functional Requirements", "20"),
        ("3.4", "Use Case Descriptions", "21"),
        ("Chapter IV", "SDLC — WATERFALL MODEL", "23"),
        ("4.1", "Software Development Life Cycle Overview", "23"),
        ("4.2", "Waterfall Model Selection Justification", "25"),
        ("4.3", "Phase 1: Requirement Analysis", "27"),
        ("4.4", "Phase 2: System Design", "29"),
        ("4.5", "Phase 3: Implementation (Coding)", "31"),
        ("4.6", "Phase 4: Testing", "33"),
        ("4.7", "Phase 5: Deployment and Maintenance", "35"),
        ("4.8", "Waterfall Model Timeline", "37"),
        ("Chapter V", "SYSTEM DESIGN", "38"),
        ("5.1", "System Architecture", "38"),
        ("5.2", "Database Design — ER Diagram", "40"),
        ("5.3", "Database Schema", "42"),
        ("5.4", "Data Flow Diagrams", "45"),
        ("5.5", "API Design", "47"),
        ("5.6", "User Interface Design", "49"),
        ("Chapter VI", "IMPLEMENTATION", "50"),
        ("6.1", "Technology Stack", "50"),
        ("6.2", "Backend Implementation (Flask)", "51"),
        ("6.3", "Database Implementation", "54"),
        ("6.4", "Frontend Implementation", "56"),
        ("6.5", "Invoice and PDF Generation", "58"),
        ("Chapter VII", "TESTING", "60"),
        ("7.1", "Testing Strategy", "60"),
        ("7.2", "Unit Testing", "61"),
        ("7.3", "Integration Testing", "63"),
        ("7.4", "User Acceptance Testing", "65"),
        ("Chapter VIII", "RESULTS AND DISCUSSION", "66"),
        ("8.1", "Dashboard and KPI Results", "66"),
        ("8.2", "Sales and Invoice Results", "68"),
        ("8.3", "Stock Management Results", "70"),
        ("8.4", "Reports Module Results", "71"),
        ("Chapter IX", "SOURCE CODE", "72"),
        ("9.1", "Backend Code (app.py)", "72"),
        ("9.2", "Frontend Code (script.js)", "76"),
        ("9.3", "Stylesheet (styles.css)", "79"),
        ("9.4", "HTML Structure (index.html)", "81"),
        ("Chapter X", "CONCLUSION AND FUTURE WORK", "83"),
        ("10.1", "Conclusion", "83"),
        ("10.2", "Limitations", "84"),
        ("10.3", "Future Scope", "85"),
        ("", "REFERENCES", "87"),
    ]
    for num, title, page in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        is_ch = num.startswith("Chapter") or (not num and title.isupper())
        if num:
            run = p.add_run(f"{num}    ")
            sf(run, FONT, BODY, bold=is_ch)
        run = p.add_run(title)
        sf(run, FONT, BODY, bold=is_ch)
        run = p.add_run(f"\t{page}")
        sf(run, FONT, BODY)
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(14), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    page_break(doc)

    # List of Figures
    heading(doc, "LIST OF FIGURES", 1)
    blank(doc, 1)
    figures = [
        ("1.1", "High-Level Block Diagram"),
        ("3.1", "Use Case Diagram — Admin Actor"),
        ("3.2", "Use Case Diagram — Viewer Actor"),
        ("4.1", "Waterfall Model Diagram"),
        ("4.2", "Waterfall Phase Timeline — Gantt Chart"),
        ("5.1", "System Architecture Diagram"),
        ("5.2", "Entity Relationship Diagram"),
        ("5.3", "DFD Level 0 — Context Diagram"),
        ("5.4", "DFD Level 1 — Process Decomposition"),
        ("5.5", "API Endpoint Flow Diagram"),
        ("6.1", "Dashboard KPI Screenshot"),
        ("6.2", "Sales Management Screenshot"),
        ("6.3", "Customer Management Screenshot"),
        ("6.4", "Products Management Screenshot"),
        ("6.5", "Stocks Module Screenshot"),
        ("6.6", "Invoice Builder Screenshot"),
        ("6.7", "Reports Module Screenshot"),
        ("6.8", "Login Authentication Screenshot"),
    ]
    for num, cap in figures:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"Figure {num} — {cap}")
        sf(run, FONT, BODY)

    page_break(doc)

    # List of Tables
    heading(doc, "LIST OF TABLES", 1)
    blank(doc, 1)
    tables = [
        ("2.1", "Comparison of Existing Systems"),
        ("3.1", "Functional Requirements List"),
        ("3.2", "Non-Functional Requirements"),
        ("4.1", "Waterfall Phase Deliverables"),
        ("4.2", "Sprint/Phase Timeline"),
        ("5.1", "Database Schema — customer_dim"),
        ("5.2", "Database Schema — product_dim"),
        ("5.3", "Database Schema — sales_fact"),
        ("5.4", "Database Schema — invoice_fact"),
        ("5.5", "Database Schema — stock_history"),
        ("5.6", "API Endpoints Summary"),
        ("6.1", "Technology Stack Summary"),
        ("7.1", "Unit Test Cases"),
        ("7.2", "Integration Test Cases"),
        ("7.3", "UAT Results Summary"),
        ("8.1", "KPI Metrics from Live System"),
        ("8.2", "Sales Records Summary"),
    ]
    for num, cap in tables:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"Table {num} — {cap}")
        sf(run, FONT, BODY)

    page_break(doc)


# ═══════════════════════════════════════════
#  NEW CHAPTER: SDLC — WATERFALL MODEL
# ═══════════════════════════════════════════

def chapter_sdlc(doc):
    chapter_title(doc, "IV", "SDLC — WATERFALL MODEL")

    heading(doc, "SOFTWARE DEVELOPMENT LIFE CYCLE", 2, "4.")

    heading(doc, "Software Development Life Cycle Overview", 2, "4.1")
    para(doc, (
        "The Software Development Life Cycle (SDLC) is a systematic "
        "process for planning, creating, testing, and deploying software "
        "applications. It provides a structured framework that defines "
        "the tasks to be performed at each stage of the software "
        "development process. The SDLC ensures that the final software "
        "product meets both functional and non-functional requirements "
        "while being delivered within time and budget constraints [17]."
    ))
    para(doc, (
        "Various SDLC models exist, each with distinct characteristics "
        "suited to different project types. The most commonly used "
        "models include:"
    ))
    bullet(doc, "Waterfall Model — Sequential, phase-based approach")
    bullet(doc, "Agile Model — Iterative, sprint-based development")
    bullet(doc, "Spiral Model — Risk-driven, iterative approach")
    bullet(doc, "V-Model — Verification and Validation model")
    bullet(doc, "Iterative Model — Repeated cycles of development")
    bullet(doc, "RAD Model — Rapid Application Development")

    para(doc, (
        "For the Retail Sales Data Analysis project, the Waterfall "
        "Model was selected as the SDLC methodology. This decision was "
        "based on the well-defined nature of the project requirements, "
        "the single-developer context, and the academic timeline "
        "constraints. The following sections describe each phase of the "
        "Waterfall Model as applied to this project."
    ))

    para(doc, (
        "The Waterfall Model, originally described by Dr. Winston W. "
        "Royce in 1970, is the earliest SDLC model and follows a "
        "linear sequential approach where each phase must be completed "
        "before the next phase begins. The output of one phase serves "
        "as the input for the subsequent phase, creating a cascading "
        "flow reminiscent of a waterfall [17]."
    ))

    para(doc, (
        "Figure 4.1 illustrates the five phases of the Waterfall Model "
        "as applied to the Retail Sales Data Analysis project."
    ))

    para(doc, "[Insert Figure 4.1: Waterfall Model Diagram]", BODY,
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         color=(150, 150, 150))
    para(doc, "Figure 4.1: Waterfall Model — Five Sequential Phases",
         SMALL, bold=True, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)

    heading(doc, "Waterfall Model Selection Justification", 2, "4.2")
    para(doc, (
        "The Waterfall Model was selected for the Retail Sales Data "
        "Analysis project based on the following justifications:"
    ))

    justifications = [
        ("Well-Defined Requirements",
         "All 24 functional requirements (Table 3.1) were identified "
         "and documented during the initial analysis phase. The project "
         "scope — dashboard, CRUD operations, invoice generation, stock "
         "management, and reports — was clearly defined before development "
         "commenced. There was no anticipated requirement volatility."),
        ("Single Developer Context",
         "The project was developed by a single MCA student. The "
         "Waterfall Model's sequential nature eliminates the overhead of "
         "sprint planning, stand-up meetings, and backlog grooming that "
         "Agile methodologies require for team coordination."),
        ("Academic Timeline",
         "The project followed a fixed academic calendar with defined "
         "submission deadlines. The Waterfall Model's milestone-based "
         "approach aligns well with academic evaluation checkpoints "
         "(mid-term review, final submission)."),
        ("Technology Stability",
         "All technologies used (Flask, PostgreSQL, HTML/CSS/JS, "
         "Chart.js, ReportLab) are mature and stable. No research or "
         "prototyping was needed to validate technology choices, "
         "eliminating the need for iterative technology evaluation."),
        ("Documentation Requirements",
         "Academic project submissions require comprehensive "
         "documentation at each stage. The Waterfall Model naturally "
         "produces documentation artifacts (SRS, design documents, test "
         "reports) at each phase transition."),
    ]
    for title, desc in justifications:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"• {title}: ")
        sf(run, FONT, BODY, bold=True)
        run = p.add_run(desc)
        sf(run, FONT, BODY)

    para(doc, (
        "While the Waterfall Model has known limitations — particularly "
        "its inflexibility to requirement changes during later phases — "
        "these limitations were mitigated by the project's well-defined "
        "scope and the single-developer's ability to make minor "
        "adjustments without formal change management processes."
    ))

    page_break(doc)

    # ── Phase 1: Requirement Analysis ──
    heading(doc, "Phase 1: Requirement Analysis", 2, "4.3")
    para(doc, (
        "The Requirement Analysis phase was the first and foundational "
        "phase of the Waterfall Model for the Retail Sales Data Analysis "
        "project. This phase focused on understanding the problem domain, "
        "identifying stakeholder needs, and documenting all functional "
        "and non-functional requirements that the system must satisfy."
    ))

    heading(doc, "Activities Performed", 3, "4.3.1")
    numbered(doc, 1, (
        "Problem Domain Study: Analysed the operational challenges "
        "faced by small and medium-sized retail businesses in India, "
        "particularly in Odisha, in managing sales data, inventory, "
        "and invoicing."
    ))
    numbered(doc, 2, (
        "Stakeholder Identification: Identified two primary actors — "
        "Admin (business owner/manager with full CRUD access) and "
        "Viewer (any user with dashboard-only access)."
    ))
    numbered(doc, 3, (
        "Requirement Elicitation: Derived 24 functional requirements "
        "(FR-01 to FR-24) and 12 non-functional requirements across "
        "performance, security, usability, maintainability, portability, "
        "and reliability categories."
    ))
    numbered(doc, 4, (
        "Feasibility Assessment: Conducted technical feasibility "
        "(open-source stack availability), economic feasibility (zero "
        "licensing cost vs. commercial alternatives), and operational "
        "feasibility (single-command startup, browser-based access) "
        "analyses."
    ))
    numbered(doc, 5, (
        "Use Case Modelling: Created use case diagrams for Admin and "
        "Viewer actors, with detailed use case description for the "
        "most complex operation (UC-12: Generate Invoice Using Invoice "
        "Builder)."
    ))

    heading(doc, "Deliverables", 3, "4.3.2")
    bullet(doc, "Software Requirements Specification (SRS) document")
    bullet(doc, "Use Case Diagrams (Admin and Viewer)")
    bullet(doc, "Feasibility Study Report")
    bullet(doc, "Functional Requirements Table (24 requirements)")
    bullet(doc, "Non-Functional Requirements Table (12 requirements)")

    heading(doc, "Duration", 3, "4.3.3")
    para(doc, (
        "The Requirement Analysis phase spanned approximately 3 weeks "
        "(Weeks 1-3 of the project timeline). This duration included "
        "literature review of existing systems, requirements elicitation "
        "through domain analysis, and documentation of the SRS."
    ))

    page_break(doc)

    # ── Phase 2: System Design ──
    heading(doc, "Phase 2: System Design", 2, "4.4")
    para(doc, (
        "The System Design phase translated the documented requirements "
        "into a detailed technical blueprint for the system. This phase "
        "produced the architectural design, database schema, API "
        "specification, and user interface wireframes that guided the "
        "implementation phase."
    ))

    heading(doc, "Activities Performed", 3, "4.4.1")
    numbered(doc, 1, (
        "Architecture Design: Designed the three-tier client-server "
        "architecture (Presentation Tier: HTML/CSS/JS, Application Tier: "
        "Flask API, Data Tier: PostgreSQL) with a security layer spanning "
        "all tiers."
    ))
    numbered(doc, 2, (
        "Database Design: Created the Entity Relationship Diagram (ERD) "
        "with 5 tables (customer_dim, product_dim, sales_fact, "
        "invoice_fact, stock_history) in Third Normal Form (3NF). Defined "
        "all primary keys, foreign keys with CASCADE constraints, data "
        "types, and default values."
    ))
    numbered(doc, 3, (
        "API Design: Specified 25+ RESTful API endpoints with HTTP "
        "methods (GET/POST/PUT/DELETE), URL patterns, authentication "
        "requirements, request/response JSON schemas, and error codes."
    ))
    numbered(doc, 4, (
        "Data Flow Diagram Design: Created DFD Level 0 (Context Diagram) "
        "showing the system boundary with Admin User and PostgreSQL as "
        "external entities, and DFD Level 1 decomposing the system into "
        "7 processes (Authenticate, Manage Customers, Manage Products, "
        "Manage Sales, Generate Invoices, Manage Stocks, Generate "
        "Reports) with 6 data stores."
    ))
    numbered(doc, 5, (
        "UI/UX Design: Designed the sidebar-main layout pattern, dark "
        "theme colour palette (15+ CSS custom properties), KPI card "
        "grid layout, form control styling, table formatting, and "
        "glassmorphism effects for the login overlay."
    ))

    heading(doc, "Deliverables", 3, "4.4.2")
    bullet(doc, "System Architecture Diagram")
    bullet(doc, "Entity Relationship Diagram (5 tables)")
    bullet(doc, "Database Schema Specification (Tables 5.1-5.5)")
    bullet(doc, "DFD Level 0 and DFD Level 1")
    bullet(doc, "API Endpoint Specification (Table 5.6)")
    bullet(doc, "UI Wireframe/Colour Palette Documentation")

    heading(doc, "Duration", 3, "4.4.3")
    para(doc, (
        "The System Design phase spanned approximately 3 weeks "
        "(Weeks 4-6). The database design and API specification were "
        "completed first, followed by the UI design and DFD creation."
    ))

    page_break(doc)

    # ── Phase 3: Implementation (Coding) ──
    heading(doc, "Phase 3: Implementation (Coding)", 2, "4.5")
    para(doc, (
        "The Implementation phase translated the design documents into "
        "working source code. This was the most effort-intensive phase, "
        "involving backend API development, database schema creation, "
        "frontend UI implementation, and PDF generation pipeline "
        "construction."
    ))

    heading(doc, "Activities Performed", 3, "4.5.1")
    numbered(doc, 1, (
        "Environment Setup: Installed Python 3.11, PostgreSQL 16, and "
        "all pip packages (flask, flask-cors, psycopg2-binary, "
        "reportlab). Created the sales_db database and configured "
        "environment variables."
    ))
    numbered(doc, 2, (
        "Backend Development (app.py — 450+ lines): Implemented all "
        "Flask route handlers for authentication (login, logout, "
        "auth_check), CRUD operations (customers, products, sales), "
        "analytics (KPIs, revenue-by-category), stock management "
        "(restock, history), invoice operations (save, download, "
        "download-custom, status update), and reports (filtered sales)."
    ))
    numbered(doc, 3, (
        "Database Schema Implementation (POST /api/setup): Coded the "
        "CREATE TABLE IF NOT EXISTS statements for all 5 tables with "
        "proper constraints, and the ALTER TABLE migration block for "
        "backward compatibility."
    ))
    numbered(doc, 4, (
        "Frontend Development (script.js — 700+ lines): Implemented "
        "the API communication layer (async api() helper), navigation "
        "system (navigateTo()), authentication flow (checkAuth, login, "
        "logout, updateAuthUI), data loading functions (loadCustomers, "
        "loadProducts, loadSales, loadKPIs, loadCharts), CRUD form "
        "handlers (submitCustomer, submitProduct, addSale), stock "
        "management (loadStocksPage, restockProduct, filterStock), "
        "invoice builder (onInvCustomerChange, renderInvRows, "
        "recalcTotals, buildInvPayload, previewInvoice, "
        "saveAndDownloadInvoice), and reports (loadReport, exportCSV)."
    ))
    numbered(doc, 5, (
        "CSS Styling (styles.css — 500+ lines): Implemented the "
        "complete dark theme with CSS custom properties, sidebar layout, "
        "KPI card grid, chart containers, form controls, table styling, "
        "badge components, button variants, login overlay with "
        "glassmorphism, toast notifications, and responsive media "
        "queries for 900px and 1100px breakpoints."
    ))
    numbered(doc, 6, (
        "HTML Structure (index.html — 350+ lines): Created the "
        "single-page application shell with sidebar navigation (7 "
        "sections), topbar with auth controls, 7 content sections "
        "(dashboard, sales, customers, products, stocks, invoices, "
        "reports), login overlay modal, invoice preview modal, and "
        "toast notification element."
    ))
    numbered(doc, 7, (
        "PDF Generation Pipeline (_build_invoice_pdf function): "
        "Implemented A4-format invoice PDF creation using ReportLab's "
        "canvas API with company header, customer billing details, "
        "itemised line items, GST calculation, terms and conditions, "
        "and branded footer."
    ))

    heading(doc, "Code Metrics", 3, "4.5.2")
    add_table(doc,
        ["File", "Language", "Lines of Code", "Functions/Routes"],
        [
            ["app.py", "Python", "~450", "25+ routes"],
            ["script.js", "JavaScript", "~700", "40+ functions"],
            ["styles.css", "CSS", "~500", "100+ rules"],
            ["index.html", "HTML", "~350", "7 sections"],
            ["TOTAL", "—", "~2,000", "—"],
        ],
        "Table 4.1: Source Code Metrics"
    )

    heading(doc, "Duration", 3, "4.5.3")
    para(doc, (
        "The Implementation phase spanned approximately 6 weeks "
        "(Weeks 7-12). Backend API development consumed 2 weeks, "
        "frontend JavaScript consumed 2 weeks, CSS styling and HTML "
        "structure consumed 1 week, and PDF generation consumed 1 week."
    ))

    page_break(doc)

    # ── Phase 4: Testing ──
    heading(doc, "Phase 4: Testing", 2, "4.6")
    para(doc, (
        "The Testing phase verified that the implemented system meets "
        "all documented requirements and functions correctly under "
        "both normal and edge-case conditions. A three-level testing "
        "strategy was employed: unit testing, integration testing, and "
        "user acceptance testing."
    ))

    heading(doc, "Activities Performed", 3, "4.6.1")
    numbered(doc, 1, (
        "Unit Testing: Tested 20 individual API endpoints (UT-01 to "
        "UT-20) covering valid inputs (happy path), invalid inputs "
        "(error path), authentication enforcement, stock validation, "
        "and PDF generation. All 20 unit tests passed."
    ))
    numbered(doc, 2, (
        "Integration Testing: Tested 6 multi-component workflows "
        "(IT-01 to IT-06) covering complete sale workflow, invoice "
        "builder auto-fill, invoice save with KPI update, mark paid "
        "with KPI update, delete sale with stock restore, and report "
        "filter with chart update. All 6 integration tests passed."
    ))
    numbered(doc, 3, (
        "User Acceptance Testing: Evaluated the system against all 24 "
        "functional requirements (FR-01 to FR-24) through systematic "
        "browser-based execution. 100% pass rate achieved across all "
        "8 modules (Authentication, Dashboard, Customers, Products, "
        "Sales, Stocks, Invoices, Reports)."
    ))

    heading(doc, "Deliverables", 3, "4.6.2")
    bullet(doc, "Unit Test Cases and Results (Table 7.1)")
    bullet(doc, "Integration Test Cases and Results (Table 7.2)")
    bullet(doc, "UAT Results Summary (Table 7.3)")
    bullet(doc, "Defect Log (0 critical defects)")

    heading(doc, "Duration", 3, "4.6.3")
    para(doc, (
        "The Testing phase spanned approximately 2 weeks (Weeks 13-14). "
        "Unit testing consumed 1 week and integration/UAT testing "
        "consumed 1 week."
    ))

    page_break(doc)

    # ── Phase 5: Deployment and Maintenance ──
    heading(doc, "Phase 5: Deployment and Maintenance", 2, "4.7")
    para(doc, (
        "The Deployment and Maintenance phase covers the operational "
        "readiness of the system and planned maintenance activities."
    ))

    heading(doc, "Deployment Activities", 3, "4.7.1")
    numbered(doc, 1, (
        "Local Deployment: The system is deployed locally using "
        "Flask's built-in Werkzeug development server on port 5000. "
        "Startup requires a single command: python app.py."
    ))
    numbered(doc, 2, (
        "Database Initialisation: The POST /api/setup endpoint creates "
        "all 5 database tables idempotently, allowing one-command "
        "database setup."
    ))
    numbered(doc, 3, (
        "Sample Data Loading: Sample data (6 customers, 8 products, "
        "6 sales) is loaded through the browser-based management "
        "interfaces."
    ))
    numbered(doc, 4, (
        "Documentation: This project report, along with inline code "
        "comments in all source files, serves as the system "
        "documentation."
    ))

    heading(doc, "Maintenance Plan", 3, "4.7.2")
    para(doc, (
        "The following maintenance activities are planned for the "
        "post-deployment period:"
    ))
    bullet(doc, (
        "Corrective Maintenance: Bug fixes identified during "
        "extended usage."
    ))
    bullet(doc, (
        "Adaptive Maintenance: Updates to accommodate new browser "
        "versions or PostgreSQL releases."
    ))
    bullet(doc, (
        "Perfective Maintenance: Performance optimisation for larger "
        "datasets (100+ products, 1000+ sales)."
    ))
    bullet(doc, (
        "Preventive Maintenance: Regular database backup procedures "
        "using pg_dump."
    ))

    heading(doc, "Duration", 3, "4.7.3")
    para(doc, (
        "Deployment consumed 1 week (Week 15). Maintenance is ongoing "
        "throughout the system's operational life."
    ))

    # ── Waterfall Timeline ──
    heading(doc, "Waterfall Model Timeline", 2, "4.8")
    para(doc, (
        "Table 4.2 presents the complete Waterfall Model timeline for "
        "the Retail Sales Data Analysis project, showing the duration, "
        "key activities, and deliverables for each phase."
    ))
    add_table(doc,
        ["Phase", "Duration", "Weeks", "Key Deliverables"],
        [
            ["Requirement Analysis", "3 weeks", "W1-W3",
             "SRS, Use Cases, Feasibility"],
            ["System Design", "3 weeks", "W4-W6",
             "Architecture, ERD, DFD, API Spec"],
            ["Implementation", "6 weeks", "W7-W12",
             "app.py, script.js, styles.css, index.html"],
            ["Testing", "2 weeks", "W13-W14",
             "Unit/Integration/UAT Test Reports"],
            ["Deployment", "1 week", "W15",
             "Deployed System, Documentation"],
            ["TOTAL", "15 weeks", "W1-W15", "Complete System + Report"],
        ],
        "Table 4.2: Waterfall Phase Timeline"
    )

    para(doc, (
        "Figure 4.2 presents a visual Gantt chart representation of "
        "the Waterfall timeline."
    ))
    para(doc, "[Insert Figure 4.2: Waterfall Phase Timeline — Gantt Chart]",
         BODY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         color=(150, 150, 150))
    para(doc, "Figure 4.2: Waterfall Phase Timeline — Gantt Chart",
         SMALL, bold=True, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break(doc)


# ═══════════════════════════════════════════
#  NEW CHAPTER: SOURCE CODE
# ═══════════════════════════════════════════

def chapter_source_code(doc):
    chapter_title(doc, "IX", "SOURCE CODE")
    para(doc, (
        "This chapter presents the complete source code of the Retail "
        "Sales Data Analysis system. The codebase consists of four files: "
        "app.py (Flask backend), script.js (frontend JavaScript), "
        "styles.css (CSS stylesheet), and index.html (HTML structure). "
        "Key sections of each file are presented with inline commentary."
    ))

    # ── 9.1 Backend Code ──
    heading(doc, "Backend Code (app.py)", 2, "9.1")
    para(doc, (
        "The Flask backend application is contained in a single file "
        "(app.py) consisting of approximately 450 lines of Python code. "
        "The following sections present the key components."
    ))

    heading(doc, "Application Configuration and Imports", 3, "9.1.1")
    code_block(doc, '''from flask import Flask, jsonify, request, send_from_directory
from flask import session, make_response
from flask_cors import CORS
from datetime import timedelta, datetime
import psycopg2, psycopg2.extras, os, functools, io

app = Flask(__name__, static_folder=".")
CORS(app, supports_credentials=True)

app.secret_key = os.getenv("SECRET_KEY", "salesdb-secret-2024")
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(minutes=30)
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"

ADMIN_USER = os.getenv("ADMIN_USER", "admin")
ADMIN_PASS = os.getenv("ADMIN_PASS", "admin123")

COMPANY = {
    "name": "SalesDB Pvt. Ltd.",
    "address": "Kendrapara, Odisha",
    "pin": "754289",
    "phone": "+91 7205109609",
    "email": "sales@salesdb.com",
    "gst": "21XXXXX1234X1ZX",
}
TAX_RATE = 0.18''', "Code 9.1: Application Configuration")

    heading(doc, "Database Connection and Authentication", 3, "9.1.2")
    code_block(doc, '''DB_CONFIG = {
    "dbname": os.getenv("DB_NAME", "sales_db"),
    "user": os.getenv("DB_USER", "postgres"),
    "password": os.getenv("DB_PASSWORD", "Abhi@4321"),
    "host": os.getenv("DB_HOST", "localhost"),
    "port": os.getenv("DB_PORT", "5432"),
    "connect_timeout": 5,
}

def get_conn():
    try:
        return psycopg2.connect(**DB_CONFIG)
    except psycopg2.OperationalError as e:
        raise RuntimeError(f"DB Error: {e}") from e

def dc(conn):
    return conn.cursor(
        cursor_factory=psycopg2.extras.RealDictCursor
    )

def login_required(f):
    @functools.wraps(f)
    def w(*a, **k):
        if not session.get("admin"):
            return jsonify({"error": "Auth required"}), 401
        return f(*a, **k)
    return w''', "Code 9.2: Database Connection and Auth Decorator")

    heading(doc, "Authentication Routes", 3, "9.1.3")
    code_block(doc, '''@app.route("/api/auth/check")
def auth_check():
    return jsonify({
        "authenticated": bool(session.get("admin"))
    })

@app.route("/api/login", methods=["POST"])
def do_login():
    d = request.get_json(force=True)
    if (d.get("username","").strip() == ADMIN_USER
        and d.get("password","") == ADMIN_PASS):
        session.permanent = True
        session["admin"] = True
        return jsonify({"success": True})
    return jsonify({
        "success": False,
        "error": "Invalid credentials"
    }), 401

@app.route("/api/logout", methods=["POST"])
def do_logout():
    session.clear()
    return jsonify({"success": True})''', "Code 9.3: Authentication Routes")

    page_break(doc)

    heading(doc, "Sales Recording Route", 3, "9.1.4")
    code_block(doc, '''@app.route("/api/sales", methods=["POST"])
@login_required
def add_sale():
    d = request.get_json(force=True)
    required = ["customer_id","product_id",
                "sale_date","quantity","sale_amount"]
    missing = [f for f in required if f not in d]
    if missing:
        return jsonify({"error":f"Missing: {missing}"}), 400
    conn = get_conn()
    with dc(conn) as cur:
        # Check stock
        cur.execute(
            "SELECT stock_qty, product_name "
            "FROM product_dim WHERE product_id=%s;",
            (d["product_id"],))
        prod = cur.fetchone()
        if not prod:
            return jsonify({"error":"Product not found"}), 404
        if (prod["stock_qty"] or 0) < d["quantity"]:
            return jsonify({
                "error": f"Insufficient stock for "
                f"{prod['product_name']}. "
                f"Available: {prod['stock_qty'] or 0}"
            }), 400
        # Insert sale
        cur.execute(
            "INSERT INTO sales_fact"
            "(customer_id,product_id,sale_date,"
            "quantity,sale_amount) "
            "VALUES(%s,%s,%s,%s,%s) RETURNING sale_id;",
            (d["customer_id"], d["product_id"],
             d["sale_date"], d["quantity"],
             d["sale_amount"]))
        sale_id = cur.fetchone()["sale_id"]
        # Deduct stock
        cur.execute(
            "UPDATE product_dim SET "
            "stock_qty=stock_qty-%s "
            "WHERE product_id=%s "
            "RETURNING stock_qty;",
            (d["quantity"], d["product_id"]))
        new_stock = cur.fetchone()["stock_qty"]
        # Log stock change
        cur.execute(
            "INSERT INTO stock_history"
            "(product_id,change_type,"
            "qty_change,new_stock) "
            "VALUES(%s,%s,%s,%s);",
            (d["product_id"], 'SALE',
             -d["quantity"], new_stock))
        # Create invoice
        sub = float(d["sale_amount"])
        tax = round(sub * TAX_RATE, 2)
        grand = round(sub + tax, 2)
        inv_no = f"INV-{datetime.now():%Y}-{sale_id:05d}"
        cur.execute(
            "INSERT INTO invoice_fact"
            "(sale_id,invoice_no,subtotal,"
            "tax_rate,tax_amount,grand_total,"
            "status) VALUES(%s,%s,%s,%s,%s,%s,%s) "
            "RETURNING invoice_id;",
            (sale_id, inv_no, sub,
             TAX_RATE*100, tax, grand, 'Pending'))
        inv_id = cur.fetchone()["invoice_id"]
    conn.commit(); conn.close()
    return jsonify({
        "sale_id": sale_id,
        "invoice_no": inv_no,
        "invoice_id": inv_id
    }), 201''', "Code 9.4: Complete Sale Recording Route")

    page_break(doc)

    heading(doc, "KPI Analytics Route", 3, "9.1.5")
    code_block(doc, '''@app.route("/api/analytics/kpis")
def kpis():
    conn = get_conn()
    with dc(conn) as cur:
        cur.execute("""
            SELECT
              COUNT(*)::int AS total_sales,
              COALESCE(SUM(sale_amount),0)::float
                AS total_revenue,
              COALESCE(AVG(sale_amount),0)::float
                AS avg_order_value,
              (SELECT COUNT(*) FROM customer_dim)::int
                AS total_customers,
              (SELECT COUNT(*) FROM product_dim)::int
                AS total_products
            FROM sales_fact;
        """)
        row = dict(cur.fetchone())
        cur.execute(
            "SELECT COUNT(*)::int AS cnt "
            "FROM product_dim "
            "WHERE COALESCE(stock_qty,0)>0 "
            "AND COALESCE(stock_qty,0)<=10;")
        row["low_stock_count"] = cur.fetchone()["cnt"]
        cur.execute(
            "SELECT COUNT(*)::int AS cnt "
            "FROM product_dim "
            "WHERE COALESCE(stock_qty,0)<=0;")
        row["out_stock_count"] = cur.fetchone()["cnt"]
        cur.execute(
            "SELECT COUNT(*)::int AS cnt "
            "FROM invoice_fact "
            "WHERE status='Pending';")
        row["pending_invoices"] = cur.fetchone()["cnt"]
    conn.close()
    return jsonify(row)''', "Code 9.5: Dashboard KPI Analytics Route")

    heading(doc, "Database Schema Setup Route", 3, "9.1.6")
    code_block(doc, '''@app.route("/api/setup", methods=["POST"])
def setup():
    conn = get_conn()
    with conn.cursor() as cur:
        cur.execute("""
          CREATE TABLE IF NOT EXISTS customer_dim(
            customer_id SERIAL PRIMARY KEY,
            first_name VARCHAR(50) NOT NULL,
            last_name VARCHAR(50) NOT NULL,
            city VARCHAR(50),
            mobile_no VARCHAR(20),
            email VARCHAR(100),
            region VARCHAR(20),
            member_type VARCHAR(20) DEFAULT 'Regular'
          );
          CREATE TABLE IF NOT EXISTS product_dim(
            product_id SERIAL PRIMARY KEY,
            product_name VARCHAR(100) NOT NULL,
            category VARCHAR(50),
            unit_price NUMERIC(10,2) DEFAULT 0,
            stock_qty INTEGER DEFAULT 0
          );
          CREATE TABLE IF NOT EXISTS sales_fact(
            sale_id SERIAL PRIMARY KEY,
            customer_id INTEGER REFERENCES
              customer_dim(customer_id)
              ON DELETE CASCADE,
            product_id INTEGER REFERENCES
              product_dim(product_id)
              ON DELETE CASCADE,
            sale_date DATE NOT NULL,
            quantity INTEGER NOT NULL,
            sale_amount NUMERIC(10,2) NOT NULL
          );
          CREATE TABLE IF NOT EXISTS invoice_fact(
            invoice_id SERIAL PRIMARY KEY,
            sale_id INTEGER REFERENCES
              sales_fact(sale_id)
              ON DELETE CASCADE,
            invoice_no VARCHAR(30) UNIQUE NOT NULL,
            subtotal NUMERIC(10,2) NOT NULL,
            tax_rate NUMERIC(5,2) DEFAULT 18,
            tax_amount NUMERIC(10,2) NOT NULL,
            grand_total NUMERIC(10,2) NOT NULL,
            status VARCHAR(20) DEFAULT 'Pending',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
          );
          CREATE TABLE IF NOT EXISTS stock_history(
            history_id SERIAL PRIMARY KEY,
            product_id INTEGER REFERENCES
              product_dim(product_id)
              ON DELETE CASCADE,
            change_type VARCHAR(20) NOT NULL,
            qty_change INTEGER NOT NULL,
            new_stock INTEGER NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
          );
        """)
    conn.commit(); conn.close()
    return jsonify({"status": "schema ready"})''',
    "Code 9.6: Database Schema Setup")

    page_break(doc)

    # ── 9.2 Frontend Code ──
    heading(doc, "Frontend Code (script.js)", 2, "9.2")
    para(doc, (
        "The frontend JavaScript code handles all client-side "
        "interactivity. Key functions are presented below."
    ))

    heading(doc, "API Communication Layer", 3, "9.2.1")
    code_block(doc, '''const API = location.protocol === 'file:'
    ? 'http://localhost:5000/api' : '/api';

async function api(path, method='GET', body=null) {
    const opts = {
        method,
        headers: {
            'Content-Type': 'application/json'
        },
        credentials: 'include'
    };
    if (body) opts.body = JSON.stringify(body);
    const res = await fetch(API + path, opts);
    if (res.status === 401) {
        isAdmin = false;
        updateAuthUI();
        navigateTo('dashboard');
        showLogin();
        toast('Session expired', 'error');
        throw new Error('Session expired');
    }
    if (!res.ok) throw new Error(await res.text());
    return res.json();
}''', "Code 9.7: API Communication Helper")

    heading(doc, "Auto-Calculation Function", 3, "9.2.2")
    code_block(doc, '''function calcSaleAmount() {
    const sel = document.getElementById('sale-product');
    const opt = sel.options[sel.selectedIndex];
    const price = parseFloat(
        opt?.getAttribute('data-price') || 0
    );
    const qty = parseInt(
        document.getElementById('sale-qty').value
    ) || 0;
    document.getElementById('sale-amount').value =
        (price * qty > 0)
            ? (price * qty).toFixed(2) : '';
}''', "Code 9.8: Auto-Calculation Function")

    heading(doc, "Dashboard Charts", 3, "9.2.3")
    code_block(doc, '''async function loadCharts() {
    const catData = await api(
        '/analytics/revenue-by-category'
    );
    if (chartCategory) chartCategory.destroy();
    if (chartDoughnut) chartDoughnut.destroy();
    if (chartPie) chartPie.destroy();

    Chart.defaults.color = '#94a3b8';
    Chart.defaults.font = {
        family: "'JetBrains Mono', monospace",
        size: 11
    };

    chartCategory = new Chart(
        document.getElementById('chartCategory'), {
        type: 'bar',
        data: {
            labels: catData.map(r => r.category),
            datasets: [{
                data: catData.map(r => r.total_revenue),
                backgroundColor: ACCENT.map(c => c+'CC'),
                borderColor: ACCENT,
                borderWidth: 1,
                borderRadius: 4,
                barThickness: 20
            }]
        },
        options: {
            responsive: true,
            maintainAspectRatio: false,
            plugins: { legend: { display: false } },
            scales: {
                x: { grid: { display: false } },
                y: {
                    grid: { color: '#334155' },
                    ticks: {
                        callback: v => '$' + v
                    }
                }
            }
        }
    });
    // ... doughnut and pie charts similarly
}''', "Code 9.9: Dashboard Chart Rendering")

    heading(doc, "Invoice Builder — Customer Selection", 3, "9.2.4")
    code_block(doc, '''async function onInvCustomerChange() {
    const cid = +document.getElementById(
        'inv-customer').value;
    const chip = document.getElementById(
        'inv-cust-chip');
    if (!cid) {
        chip.style.display = 'none';
        invRows = [];
        renderInvRows();
        recalcTotals();
        return;
    }
    const c = customers.find(
        x => x.customer_id === cid);
    if (c) {
        document.getElementById(
            'inv-cust-name').textContent =
            `${c.first_name} ${c.last_name}`;
        chip.style.display = 'block';
    }
    try {
        const salesData = await api(
            `/customers/${cid}/sales`);
        if (salesData && salesData.length) {
            invRows = salesData.map(s => ({
                sale_id: s.sale_id,
                description: s.description,
                unit_price: parseFloat(
                    s.unit_price || 0),
                qty: parseInt(s.qty || 1),
                amount: parseFloat(s.amount || 0),
            }));
            toast(`Loaded ${invRows.length} sales`);
        } else {
            invRows = [];
        }
    } catch (e) {
        invRows = [];
    }
    renderInvRows();
    recalcTotals();
}''', "Code 9.10: Invoice Builder Customer Selection")

    page_break(doc)

    # ── 9.3 CSS ──
    heading(doc, "Stylesheet (styles.css)", 2, "9.3")
    para(doc, (
        "The CSS stylesheet defines the complete dark theme with CSS "
        "custom properties, responsive layout, and component styling. "
        "Key sections are presented below."
    ))

    heading(doc, "CSS Custom Properties (Design Tokens)", 3, "9.3.1")
    code_block(doc, ''':root {
    --bg: #0f172a;
    --surface: #1e293b;
    --surface2: #334155;
    --border: #334155;
    --accent: #ec4899;
    --accent2: #3b82f6;
    --accent3: #f43f5e;
    --text: #f8fafc;
    --muted: #94a3b8;
    --font-head: 'Syne', sans-serif;
    --font-mono: 'JetBrains Mono', monospace;
    --radius: 12px;
    --transition: 180ms ease;
}

body {
    background: radial-gradient(
        circle at top right,
        #1e1b4b, var(--bg)
    );
    color: var(--text);
    font-family: var(--font-mono);
    font-size: 13px;
    min-height: 100vh;
    display: flex;
}''', "Code 9.11: CSS Custom Properties")

    heading(doc, "Sidebar and KPI Card Styles", 3, "9.3.2")
    code_block(doc, '''.sidebar {
    width: 240px;
    min-height: 100vh;
    background: rgba(30, 41, 59, 0.7);
    backdrop-filter: blur(10px);
    border-right: 1px solid var(--border);
    display: flex;
    flex-direction: column;
    position: sticky;
    top: 0;
    height: 100vh;
}

.kpi-card {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 18px;
    position: relative;
    overflow: hidden;
    transition: var(--transition);
}

.kpi-card:hover {
    transform: translateY(-2px);
    border-color: var(--accent);
}

.kpi-card::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 3px;
    background: var(--accent);
}''', "Code 9.12: Sidebar and KPI Card Styles")

    page_break(doc)

    heading(doc, "Login Overlay Styles", 3, "9.3.3")
    code_block(doc, '''.login-overlay {
    position: fixed;
    top: 0; left: 0;
    right: 0; bottom: 0;
    background: rgba(15, 23, 42, 0.92);
    backdrop-filter: blur(12px);
    display: flex;
    align-items: center;
    justify-content: center;
    z-index: 9000;
    opacity: 0;
    visibility: hidden;
    transition: all 300ms ease;
}
.login-overlay.visible {
    opacity: 1;
    visibility: visible;
}
.login-card {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 16px;
    padding: 40px 34px;
    width: 370px;
    max-width: 90vw;
    box-shadow: 0 25px 50px rgba(0,0,0,0.5);
}
.login-card::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 4px;
    background: linear-gradient(
        90deg, #fbbf24, #ec4899, #3b82f6
    );
}''', "Code 9.13: Login Overlay Glassmorphism Styles")

    # ── 9.4 HTML ──
    heading(doc, "HTML Structure (index.html)", 2, "9.4")
    para(doc, (
        "The HTML file defines the single-page application structure "
        "with sidebar navigation, 7 content sections, login overlay, "
        "and invoice preview modal. Key structural elements are shown."
    ))

    heading(doc, "Sidebar Navigation", 3, "9.4.1")
    code_block(doc, '''<aside class="sidebar">
  <div class="logo">
    <div class="logo-mark">
      Sales<span>DB</span>
    </div>
    <div class="logo-sub">Analytics</div>
  </div>
  <nav class="nav">
    <div class="nav-label">Overview</div>
    <div class="nav-item active"
         data-section="dashboard">
      <span class="nav-icon">&#9638;</span>
      <span>Dashboard</span>
    </div>
    <div class="nav-item"
         data-section="sales"
         data-protected="true">
      <span class="nav-icon">&#9672;</span>
      <span>Sales</span>
      <span class="lock-icon"
            id="lock-sales">&#128274;</span>
    </div>
    <!-- Customers, Products, Stocks,
         Invoices, Reports similarly -->
  </nav>
  <div class="status-bar">
    <span class="status-dot"></span>
    <span>API Connected</span>
  </div>
</aside>''', "Code 9.14: Sidebar Navigation HTML")

    heading(doc, "Dashboard KPI Cards", 3, "9.4.2")
    code_block(doc, '''<section class="section active"
         id="section-dashboard">
  <div class="kpi-grid">
    <div class="kpi-card">
      <div class="kpi-label">Total Revenue</div>
      <div class="kpi-value"
           id="kpi-revenue">&#8212;</div>
    </div>
    <div class="kpi-card">
      <div class="kpi-label">Total Sales</div>
      <div class="kpi-value"
           id="kpi-sales">&#8212;</div>
    </div>
    <!-- 6 more KPI cards -->
  </div>
  <div class="charts-row">
    <div class="chart-card">
      <div class="chart-title">
        Revenue by Category
      </div>
      <div class="chart-wrap">
        <canvas id="chartCategory"></canvas>
      </div>
    </div>
    <!-- 2 more chart cards -->
  </div>
</section>''', "Code 9.15: Dashboard KPI Cards HTML")

    page_break(doc)


# ═══════════════════════════════════════════
#  REMAINING CHAPTERS (Existing content preserved)
# ═══════════════════════════════════════════

# NOTE: Chapters I, II, III, V (Design), VI (Implementation),
# VII (Testing), VIII (Results) use the SAME content from
# your existing PDF. I'm including the chapter stubs here.
# You should copy the content from your existing document.

def chapter_intro(doc):
    """Chapter I - same as existing pages 9-12"""
    chapter_title(doc, "I", "INTRODUCTION")
    heading(doc, "INTRODUCTION", 2, "1.")
    heading(doc, "Background and Motivation", 2, "1.1")
    para(doc, (
        "In the contemporary era of digital transformation, data-driven "
        "decision-making has become the cornerstone of competitive "
        "advantage in the retail sector. Small and medium-sized retail "
        "enterprises generate substantial volumes of transactional data "
        "daily, encompassing customer purchases, product inventory "
        "movements, payment collections, and tax obligations. However, "
        "the absence of integrated analytical systems means that this "
        "data remains largely underutilised, leaving business owners "
        "without the insights necessary to optimise pricing, manage "
        "inventory, and forecast revenue [1]."
    ))
    para(doc, (
        "The motivation for developing the Retail Sales Data Analysis "
        "system emerges from this observable gap. Traditional approaches "
        "to retail sales management in Indian SMEs typically involve a "
        "combination of manual ledgers, Microsoft Excel spreadsheets, or "
        "standalone billing software that operate in isolation. These "
        "approaches are characterised by manual data entry errors, lack "
        "of real-time inventory visibility, absence of analytical "
        "reporting, and no automated invoice generation — all of which "
        "compound to result in inefficient operations and missed business "
        "opportunities [2]."
    ))
    para(doc, (
        "Modern open-source web technologies have dramatically lowered "
        "the cost and complexity of building integrated business "
        "intelligence systems. The Python Flask framework provides a "
        "lightweight yet powerful foundation for RESTful API development. "
        "PostgreSQL offers enterprise-grade relational database "
        "capabilities at zero licensing cost. HTML5, CSS3, and JavaScript "
        "enable rich interactive frontend experiences without framework "
        "complexity. Chart.js delivers professional-quality data "
        "visualisations directly in the browser. The combination of these "
        "technologies makes it feasible to build a production-quality "
        "retail analytics system that can be deployed on any standard "
        "machine with minimal setup overhead [3]."
    ))
    para(doc, (
        "The SalesDB Analytics System — the implementation platform for "
        "this project — was designed to demonstrate how a carefully "
        "architected full-stack web application can transform raw sales "
        "transactional data into actionable business intelligence. The "
        "system integrates customer relationship management, product "
        "inventory control, automated GST-compliant invoice generation, "
        "and multi-dimensional analytics reporting into a single unified "
        "dashboard interface, accessible through any modern web browser "
        "at localhost:5000."
    ))
    para(doc, (
        "The project also serves an academic purpose of demonstrating "
        "the application of core computer science principles in a "
        "real-world context: database normalisation theory in schema "
        "design, REST architectural constraints in API design, "
        "client-server architecture in system deployment, and software "
        "engineering best practices in code organisation and error "
        "handling."
    ))

    # 1.2 Problem Statement
    heading(doc, "Problem Statement", 2, "1.2")
    para(doc, (
        "The specific problem addressed by this project can be stated "
        "as follows: Small and medium-sized retail businesses in India "
        "lack an affordable, integrated, and analytically capable sales "
        "management system that can automate their complete sales "
        "pipeline from customer registration through product sale, stock "
        "management, invoice generation, and revenue reporting, within a "
        "single web-accessible platform."
    ))
    para(doc, "This problem manifests through the following observable symptoms:")
    bullet(doc, (
        "Customer data is maintained separately from sales data, making "
        "it impossible to quickly identify a customer's purchase history, "
        "outstanding invoices, or total lifetime value [1]."
    ))
    bullet(doc, (
        "Product stock levels are not automatically decremented when "
        "sales are recorded, leading to situations where sales are "
        "accepted for out-of-stock products."
    ))
    bullet(doc, (
        "Invoice generation is performed manually using word processors "
        "or printed templates, consuming significant administrative time "
        "and introducing errors in GST calculation."
    ))
    bullet(doc, (
        "Management has no real-time visibility into key performance "
        "indicators such as total revenue, revenue by product category, "
        "pending payment collections, or average order value."
    ))
    bullet(doc, (
        "Generating filtered sales reports for specific time periods, "
        "product categories, or geographic regions requires manual "
        "extraction and manipulation of spreadsheet data [2]."
    ))

    # 1.3 Objectives
    heading(doc, "Objectives of the Project", 2, "1.3")
    objs = [
        "To design a fully normalised (3NF) relational database schema "
        "using PostgreSQL that integrates customer, product, sales, "
        "invoice, and stock history data.",
        "To develop a comprehensive RESTful API using Flask (Python) "
        "with 20+ endpoints covering all business operations.",
        "To build an interactive, dark-themed, responsive frontend "
        "dashboard using HTML5, CSS3, and JavaScript with Chart.js.",
        "To implement automated GST-inclusive (18%) invoice generation "
        "with automatic and standalone Invoice Builder modes.",
        "To develop PDF invoice download capability using ReportLab.",
        "To implement a stock management module with real-time tracking, "
        "restocking, and history logging.",
        "To build filterable sales reports with CSV export capability.",
        "To secure all write operations behind session-based admin "
        "authentication with 30-minute automatic expiry.",
    ]
    for i, obj in enumerate(objs):
        numbered(doc, i + 1, obj)

    # 1.4 Scope
    heading(doc, "Scope of the Project", 2, "1.4")
    para(doc, "The system covers the following functional modules:")
    modules = [
        "Dashboard Module: 8 KPI metrics and 3 Chart.js visualisations.",
        "Sales Module: Recording with auto-calculation and stock deduction.",
        "Customer Management: CRUD with 7 fields and membership types.",
        "Product Management: CRUD with pricing and stock quantities.",
        "Stock Management: Real-time tracking with restock and history.",
        "Invoice Module: Builder with auto-fill, GST, preview, PDF.",
        "Reports Module: Filtered analytics with CSV export.",
    ]
    for m in modules:
        bullet(doc, m)

    para(doc, (
        "Out of scope: multi-user roles, cloud deployment, mobile app, "
        "barcode scanning, email delivery. These are future enhancements."
    ))

    # 1.5 Organisation
    heading(doc, "Organisation of the Report", 2, "1.5")
    para(doc, (
        "This project report is organised into ten chapters. Chapter I "
        "introduces the project. Chapter II presents the literature "
        "review. Chapter III covers system analysis and requirements. "
        "Chapter IV details the SDLC process using the Waterfall Model. "
        "Chapter V presents system design. Chapter VI covers "
        "implementation. Chapter VII documents testing. Chapter VIII "
        "presents results and discussion. Chapter IX contains the "
        "complete source code. Chapter X presents conclusion and future "
        "work, followed by references."
    ))

    para(doc, "[Insert Figure 1.1: High-Level Block Diagram]", BODY,
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         color=(150, 150, 150))

    page_break(doc)


def chapter_conclusion(doc):
    """Chapter X - Conclusion and Future Work"""
    chapter_title(doc, "X", "CONCLUSION AND FUTURE WORK")

    heading(doc, "CONCLUSION AND FUTURE WORK", 2, "10.")

    heading(doc, "Conclusion", 2, "10.1")
    para(doc, (
        "The Retail Sales Data Analysis system has been successfully "
        "designed, implemented, tested, and evaluated as the major "
        "project for the degree of Master in Computer Application at "
        f"{COLLEGE}, Bhubaneswar. The system fulfils all 24 functional "
        "requirements identified during the requirements analysis phase, "
        "achieving 100% functional requirement coverage as verified "
        "through the layered testing strategy documented in Chapter VII."
    ))
    para(doc, (
        "The project demonstrates that a full-stack web application "
        "built on the open-source technology stack of Flask + PostgreSQL "
        "+ HTML/CSS/JavaScript + Chart.js + ReportLab can deliver "
        "enterprise-grade retail sales analytics capabilities at zero "
        "licensing cost. The system integrates customer relationship "
        "management, product inventory control, automated invoice "
        "generation, stock change tracking, and multi-dimensional "
        "analytics reporting into a single, cohesive, browser-accessible "
        "platform."
    ))
    para(doc, (
        "The project followed the Waterfall Model SDLC methodology, "
        "progressing through five sequential phases over 15 weeks: "
        "Requirement Analysis (3 weeks), System Design (3 weeks), "
        "Implementation (6 weeks), Testing (2 weeks), and Deployment "
        "(1 week). All phase deliverables were completed on schedule."
    ))
    para(doc, (
        "From the academic perspective, the project successfully applies "
        "database normalisation (3NF), RESTful architecture, MVC "
        "separation of concerns, ACID transaction compliance, and "
        "event-driven programming in a real-world context."
    ))
    para(doc, (
        "The system produces actionable business intelligence: "
        "identifying Electronics as the dominant revenue category "
        "(52.8%), flagging 62.5% of products as low or out of stock, "
        "revealing pending invoice collections of $1,157.40, and "
        "providing customer-level transaction history."
    ))

    heading(doc, "Limitations", 2, "10.2")
    limitations = [
        "Single-User Architecture with no role-based access control.",
        "Local deployment only — no cloud or containerisation.",
        "Single-page invoice PDF — no multi-page support.",
        "No automated testing framework (pytest/Newman).",
        "No email integration for invoice delivery.",
        "Currency limited to USD notation.",
        "No automated database backup mechanism.",
    ]
    for lim in limitations:
        bullet(doc, lim)

    heading(doc, "Future Scope", 2, "10.3")
    enhancements = [
        ("Role-Based Access Control",
         "Implement users table with admin/manager/cashier/viewer roles."),
        ("Cloud Deployment",
         "Containerise with Docker, deploy to AWS/Heroku with Nginx + "
         "Gunicorn and SSL encryption."),
        ("Mobile-Responsive Enhancement",
         "Full mobile support with CSS responsive breakpoints."),
        ("Barcode/QR Scanning",
         "Browser-based barcode scanning for product identification."),
        ("Email Invoice Delivery",
         "Flask-Mail or SendGrid integration for auto-emailing PDFs."),
        ("Multi-Currency Support",
         "INR/USD/EUR with Indian numbering system support."),
        ("Advanced Analytics",
         "Time-series forecasting, customer lifetime value, product "
         "affinity analysis."),
        ("Automated Backup",
         "Scheduled pg_dump with cloud storage sync."),
    ]
    for i, (title, desc) in enumerate(enhancements):
        numbered(doc, i + 1, f"{title}: {desc}")

    page_break(doc)


def references(doc):
    heading(doc, "REFERENCES", 1)
    blank(doc, 1)
    refs = [
        "Chaudhuri, S., Dayal, U., & Narasayya, V. (2011). An overview "
        "of business intelligence technology. Communications of the ACM, "
        "54(8), 88-98.",
        "Xu, L. D., & Duan, L. (2019). Big data for cyber physical "
        "systems in industry 4.0. Enterprise Information Systems, "
        "13(2), 148-169.",
        "Grinberg, M. (2018). Flask Web Development (2nd ed.). O'Reilly.",
        "Fielding, R. T. (2000). Architectural Styles and the Design of "
        "Network-based Software Architectures. UC Irvine.",
        "Greenberg, P. (2009). CRM at the Speed of Light (4th ed.). "
        "McGraw-Hill.",
        "Garrett, J. J. (2005). Ajax: A New Approach to Web Applications.",
        "Kiran, M. et al. (2015). Lambda architecture for big data "
        "processing. IEEE Big Data Conference, 2785-2792.",
        "Ramakrishnan, R., & Gehrke, J. (2003). Database Management "
        "Systems (3rd ed.). McGraw-Hill.",
        "Laudon, K. C. (2020). Management Information Systems (16th ed.). "
        "Pearson.",
        "Singh, A. & Malhotra, M. (2020). Comparative study of Python "
        "web frameworks. IJCA, 175(12), 35-42.",
        "Stonebraker, M. (2005). What goes around comes around. Readings "
        "in Database Systems. MIT Press.",
        "Pallets Projects. (2023). Flask Documentation (3.0.x). "
        "https://flask.palletsprojects.com/",
        "Di Gregorio, F. (2010). psycopg2 Documentation. "
        "https://www.psycopg.org/docs/",
        "ReportLab Inc. (2023). ReportLab User Guide. "
        "https://www.reportlab.com/docs/",
        "Chart.js Contributors. (2023). Chart.js Documentation v4.4.1. "
        "https://www.chartjs.org/docs/",
        "Lutz, M. (2013). Learning Python (5th ed.). O'Reilly.",
        "Fowler, M. (2002). Patterns of Enterprise Application "
        "Architecture. Addison-Wesley.",
        "PostgreSQL Documentation. (2023). https://www.postgresql.org/docs/",
        "Scambray, J. (2010). Hacking Exposed Web Applications (3rd ed.).",
        "AWS Documentation. (2023). AWS Elastic Beanstalk Developer Guide.",
        "Royce, W. W. (1970). Managing the Development of Large Software "
        "Systems. IEEE WESCON, 1-9.",
    ]
    for i, ref in enumerate(refs):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.hanging_indent = Cm(1)
        run = p.add_run(f"[{i + 1}]  ")
        sf(run, FONT, BODY, bold=True)
        run = p.add_run(ref)
        sf(run, FONT, BODY)


# ═══════════════════════════════════════════
#  STUB FUNCTIONS FOR EXISTING CHAPTERS
#  (Copy content from your existing PDF)
# ═══════════════════════════════════════════

def chapter_lit_review(doc):
    """Chapter II - Copy from existing PDF pages 13-16"""
    chapter_title(doc, "II", "LITERATURE REVIEW")
    heading(doc, "LITERATURE REVIEW", 2, "2.")
    heading(doc, "Overview of Web-Based Sales and Analytics Systems", 2, "2.1")
    para(doc, "[Copy content from your existing Chapter II, Section 2.1 — approximately 2 pages of text about web-based sales systems evolution, three generations, etc.]")
    heading(doc, "Related Work and Existing Systems", 2, "2.2")
    para(doc, "[Copy the comparison table and analysis from Section 2.2]")
    heading(doc, "Technologies Reviewed", 2, "2.3")
    para(doc, "[Copy Flask, PostgreSQL, ReportLab, Chart.js reviews from Section 2.3]")
    heading(doc, "Summary and Research Gap", 2, "2.4")
    para(doc, "[Copy research gap summary from Section 2.4]")
    page_break(doc)


def chapter_sys_analysis(doc):
    """Chapter III - Copy from existing PDF pages 17-22"""
    chapter_title(doc, "III", "SYSTEM ANALYSIS AND REQUIREMENTS")
    heading(doc, "SYSTEM ANALYSIS AND REQUIREMENTS", 2, "3.")
    heading(doc, "Feasibility Study", 2, "3.1")
    para(doc, "[Copy feasibility study from existing Chapter III]")
    heading(doc, "Functional Requirements", 2, "3.2")
    para(doc, "[Copy the FR table from existing document]")
    heading(doc, "Non-Functional Requirements", 2, "3.3")
    para(doc, "[Copy NFR table from existing document]")
    heading(doc, "Use Case Descriptions", 2, "3.4")
    para(doc, "[Copy use case diagrams and descriptions]")
    page_break(doc)


def chapter_design(doc):
    """Chapter V (now) - Copy from existing PDF pages 23-29"""
    chapter_title(doc, "V", "SYSTEM DESIGN")
    heading(doc, "SYSTEM DESIGN", 2, "5.")
    heading(doc, "System Architecture", 2, "5.1")
    para(doc, "[Copy architecture content from existing Chapter IV]")
    heading(doc, "Database Design — ER Diagram", 2, "5.2")
    para(doc, "[Copy ER diagram and explanation]")
    heading(doc, "Database Schema", 2, "5.3")
    para(doc, "[Copy all 5 table schema tables]")
    heading(doc, "Data Flow Diagrams", 2, "5.4")
    para(doc, "[Copy DFD Level 0 and Level 1 diagrams and descriptions]")
    heading(doc, "API Design", 2, "5.5")
    para(doc, "[Copy API endpoints table]")
    heading(doc, "User Interface Design", 2, "5.6")
    para(doc, "[Copy UI design description]")
    page_break(doc)


def chapter_implementation(doc):
    """Chapter VI - Copy from existing PDF pages 30-37"""
    chapter_title(doc, "VI", "IMPLEMENTATION")
    heading(doc, "IMPLEMENTATION", 2, "6.")
    heading(doc, "Technology Stack", 2, "6.1")
    para(doc, "[Copy tech stack table from existing Chapter V]")
    heading(doc, "Backend Implementation (Flask)", 2, "6.2")
    para(doc, "[Copy backend implementation details]")
    heading(doc, "Database Implementation (PostgreSQL)", 2, "6.3")
    para(doc, "[Copy database implementation details]")
    heading(doc, "Frontend Implementation", 2, "6.4")
    para(doc, "[Copy frontend implementation details + screenshots]")
    heading(doc, "Invoice and PDF Generation", 2, "6.5")
    para(doc, "[Copy invoice/PDF generation details]")
    page_break(doc)


def chapter_testing(doc):
    """Chapter VII - Copy from existing PDF pages 38-41"""
    chapter_title(doc, "VII", "TESTING")
    heading(doc, "TESTING", 2, "7.")
    heading(doc, "Testing Strategy", 2, "7.1")
    para(doc, "[Copy testing strategy from existing Chapter VI]")
    heading(doc, "Unit Testing", 2, "7.2")
    para(doc, "[Copy unit test table]")
    heading(doc, "Integration Testing", 2, "7.3")
    para(doc, "[Copy integration test table]")
    heading(doc, "User Acceptance Testing", 2, "7.4")
    para(doc, "[Copy UAT results table]")
    page_break(doc)


def chapter_results(doc):
    """Chapter VIII - Copy from existing PDF pages 42-44"""
    chapter_title(doc, "VIII", "RESULTS AND DISCUSSION")
    heading(doc, "RESULTS AND DISCUSSION", 2, "8.")
    heading(doc, "Dashboard and KPI Results", 2, "8.1")
    para(doc, "[Copy KPI results table and analysis]")
    heading(doc, "Sales and Invoice Results", 2, "8.2")
    para(doc, "[Copy sales records table and analysis]")
    heading(doc, "Stock Management Results", 2, "8.3")
    para(doc, "[Copy stock management analysis]")
    heading(doc, "Reports Module Results", 2, "8.4")
    para(doc, "[Copy reports analysis]")
    page_break(doc)


# ═══════════════════════════════════════════
#  MAIN BUILD
# ═══════════════════════════════════════════

def build():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(BODY)

    # Front matter
    cover_page(doc)
    certificate(doc)
    acknowledgement(doc)
    declaration(doc)
    abstract(doc)
    table_of_contents(doc)

    # Chapters
    chapter_intro(doc)           # Ch I  (~4 pages)
    chapter_lit_review(doc)      # Ch II (~4 pages) - COPY FROM PDF
    chapter_sys_analysis(doc)    # Ch III (~6 pages) - COPY FROM PDF
    chapter_sdlc(doc)            # Ch IV — NEW SDLC CHAPTER (~10 pages)
    chapter_design(doc)          # Ch V  (~7 pages) - COPY FROM PDF
    chapter_implementation(doc)  # Ch VI (~8 pages) - COPY FROM PDF
    chapter_testing(doc)         # Ch VII (~4 pages) - COPY FROM PDF
    chapter_results(doc)         # Ch VIII (~3 pages) - COPY FROM PDF
    chapter_source_code(doc)     # Ch IX — NEW SOURCE CODE (~12 pages)
    chapter_conclusion(doc)      # Ch X  (~3 pages)
    references(doc)              # (~2 pages)

    header_footer(doc)

    fname = "Retail_Sales_Data_Analysis_Final.docx"
    doc.save(fname)
    return fname


if __name__ == '__main__':
    print('\n  ╔════════════════════════════════════════════╗')
    print('  ║  Retail Sales Data Analysis — Final Report  ║')
    print('  ║  Target: 75-85 pages                        ║')
    print('  ║  Font: Arial (14/12/11 pt)                  ║')
    print('  ╚════════════════════════════════════════════╝\n')

    fname = build()

    print(f'  ✅ Saved: {fname}')
    print('  ✅ NEW additions:')
    print('     • Chapter IV: SDLC — Waterfall Model (~10 pages)')
    print('     • Chapter IX: Source Code (~12 pages)')
    print('     • Updated Table of Contents')
    print('     • Updated Abstract with SDLC mention')
    print('  ✅ Chapters with [Copy...] placeholders:')
    print('     Replace with content from your existing PDF\n')

    chapters = [
        'Cover Page + Certificate + Declaration + Acknowledgement (4 pg)',
        'Abstract (1 pg)',
        'Table of Contents + Lists (3 pg)',
        'Ch I   — Introduction (4 pg)',
        'Ch II  — Literature Review (4 pg) [COPY FROM PDF]',
        'Ch III — System Analysis (6 pg) [COPY FROM PDF]',
        'Ch IV  — SDLC Waterfall Model (10 pg) ★NEW',
        'Ch V   — System Design (7 pg) [COPY FROM PDF]',
        'Ch VI  — Implementation (8 pg) [COPY FROM PDF]',
        'Ch VII — Testing (4 pg) [COPY FROM PDF]',
        'Ch VIII— Results (3 pg) [COPY FROM PDF]',
        'Ch IX  — Source Code (12 pg) ★NEW',
        'Ch X   — Conclusion & Future (3 pg)',
        'References (2 pg)',
    ]
    total = 4+1+3+4+4+6+10+7+8+4+3+12+3+2
    print(f'  Estimated pages: ~{total}')
    print('  Chapters:')
    for c in chapters:
        print(f'    ✓ {c}')
    print()